import logging
import re
import asyncio
from typing import Optional, Tuple
from core.agents.state import TargetedEditingState
from core.agents.section_editor import create_section_editing_graph

logger = logging.getLogger(__name__)


def parse_example_node(state: TargetedEditingState) -> dict:
    """
    Node 1: Parse example document into sections.
    """
    logger.info("Parsing example document into sections")
    
    sections = parse_document_structure(state.example_document_text)
    
    # Enhance with metadata
    enhanced = {}
    for key, content in sections.items():
        enhanced[key] = {
            "key": key,
            "title": denormalize_section_name(key),
            "content": content
        }
    
    logger.info(f"Found {len(enhanced)} sections in example document")
    
    return {"example_sections": enhanced}


async def edit_single_section(
    section_change,
    example_sections: dict,
    example_document_text: str,
    sources: list[str],
    graph
) -> Tuple[Optional[str], Optional[dict]]:
    """
    Edit a single section using the provided graph.
    
    Args:
        section_change: SectionChange object with section_name and user_direction
        example_sections: Dictionary of parsed sections from example document
        example_document_text: Full text of example document
        sources: List of source filenames for RAG retrieval
        graph: Compiled section editing graph (reused for efficiency)
    
    Returns:
        Tuple of (matching_key, section_data) or (None, None) if section not found
    """
    # Normalize section name for matching
    section_key = normalize_section_name(section_change.section_name)
    
    # Find matching section (fuzzy match)
    matching_key = None
    for key in example_sections.keys():
        if section_key in key or key in section_key:
            matching_key = key
            break
    
    if not matching_key:
        logger.warning(f"Section '{section_change.section_name}' not found in example, skipping")
        return None, None
    
    section_data = example_sections[matching_key]
    
    logger.info(f"Editing section: {section_data['title']}")
    
    # Edit using the shared graph (async invoke for parallel execution)
    result = await graph.ainvoke({
        "section_title": section_data["title"],
        "original_content": section_data["content"],
        "user_direction": section_change.user_direction,
        "full_document": example_document_text,
        "sources": sources
    })
    
    logger.info(f"  ✓ Section edited: {section_data['title']} ({len(result['new_content'])} chars)")
    
    return matching_key, {
        "title": result["new_title"],
        "content": result["new_content"]
    }


async def edit_sections_node(state: TargetedEditingState) -> dict:
    """
    Node 2: Edit specified sections based on user directions.
    Uses parallel execution for better performance.
    """
    logger.info(f"Editing {len(state.section_changes)} sections")
    
    # Create graph ONCE (all sections use same sources, so reuse for efficiency)
    sources = list(state.reference_texts.keys())
    section_edit_graph = create_section_editing_graph(sources=sources)
    
    # Prepare all edit tasks for parallel execution
    edit_tasks = [
        edit_single_section(
            section_change,
            state.example_sections,
            state.example_document_text,
            sources,
            section_edit_graph
        )
        for section_change in state.section_changes
    ]
    
    # Execute all edits in parallel using asyncio.gather
    logger.info(f"Running {len(edit_tasks)} section edits in parallel...")
    results = await asyncio.gather(*edit_tasks)
    
    # Collect results (filter out sections that weren't found)
    modified_sections = {}
    for matching_key, section_data in results:
        if matching_key is not None:
            modified_sections[matching_key] = section_data
    
    logger.info(f"Completed editing {len(modified_sections)} sections")
    
    return {"modified_sections": modified_sections}


def assemble_document_node(state: TargetedEditingState) -> dict:
    """
    Node 3: Assemble final document with modified and unchanged sections.
    """
    logger.info("Assembling final document")
    
    from docx import Document
    import os
    
    doc = Document()
    modified_count = 0
    unchanged_count = 0
    
    for section_key, section_data in state.example_sections.items(): 
        if section_key in state.modified_sections:
            doc.add_heading(state.modified_sections[section_key]["title"], level=1)
            content = state.modified_sections[section_key]["content"]
            modified_count += 1
            logger.info(f" {state.modified_sections[section_key]['title']}: MODIFIED")
        else:
            doc.add_heading(section_data["title"], level=1)
            content = section_data["content"]
            unchanged_count += 1
            logger.info(f"  {section_data['title']}: unchanged")
        
        # Add paragraphs
        for para in content.split('\n\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
    
    # Save document
    os.makedirs(os.path.dirname(state.output_filename) or 'outputs', exist_ok=True)
    doc.save(state.output_filename)
    
    logger.info(f"Document saved: {modified_count} modified, {unchanged_count} unchanged")
    
    return {
        "stats": {
            "total_sections": len(state.example_sections),
            "modified": modified_count,
            "unchanged": unchanged_count
        }
    }


# ============================================================================
# DOCUMENT PARSING UTILITIES
# ============================================================================

def parse_document_structure(text: str) -> dict[str, str]:
    """
    Parse document into sections based on headings.
    """
    sections = {}
    current_section = None
    current_content = []
    
    lines = text.split('\n')
    
    # Debug: log first 20 lines to understand document structure
    logger.debug(f"Document has {len(lines)} lines")
    logger.debug("First 20 lines of document:")
    for i, line in enumerate(lines[:20]):
        is_head = is_heading(line)
        logger.debug(f"  Line {i}: [{'HEADING' if is_head else 'content'}] {line[:80]}")
    
    for line in lines:
        if is_heading(line):
            # Save previous section
            if current_section:
                content = '\n'.join(current_content).strip()
                key = normalize_section_name(current_section)
                sections[key] = content
                logger.debug(f"Saved section: '{current_section}' -> key='{key}', length={len(content)}")
            
            # Start new section
            current_section = line.strip().rstrip(':.')
            current_content = []
            logger.debug(f"New section started: '{current_section}'")
        else:
            if current_section:
                current_content.append(line)
    
    # Save last section
    if current_section:
        content = '\n'.join(current_content).strip()
        key = normalize_section_name(current_section)
        sections[key] = content
        logger.debug(f"Saved final section: '{current_section}' -> key='{key}', length={len(content)}")
    
    # If no sections found, treat entire document as one section
    if not sections:
        logger.warning("No sections detected! Treating entire document as 'Main Content'")
        sections["main_content"] = text.strip()
    
    return sections


def is_heading(line: str) -> bool:
    """
    Detect if a line is a heading - balanced approach.
    Matches clear heading patterns while avoiding common false positives.
    """
    line = line.strip()
    
    if not line:
        return False
    
    # Skip very long lines (likely not headings)
    if len(line) > 100:
        return False
    
    words = line.split()
    word_count = len(words)
    
    # Pattern 1: All caps (but not too long)
    if line.isupper() and 3 < len(line) < 80:
        return True
    
    # Pattern 2: Numbered headings (e.g., "1. Introduction" or "1.1 Background")
    if re.match(r'^\d+(\.\d+)*\.?\s+[A-Z]', line):
        return True
    
    # Pattern 3: Special headings starting with How/What/Why/When/Where (procedural headings)
    # Check this FIRST before capitalization, as these often have low cap ratios
    if 2 <= word_count <= 12 and not line.endswith(('.', ',', ';', '!', '?')):
        if words[0].lower() in ['how', 'what', 'why', 'when', 'where', 'who']:
            # "How to..." style headings - allow more common words
            lower_words = [w.lower() for w in words]
            sentence_words = {'the', 'a', 'an', 'is', 'are', 'was', 'were', 'to', 'of', 'in', 'on', 'for', 'with', 'at', 'by', 'and'}
            common_count = len(sentence_words & set(lower_words))
            # If not too many common words (< half the line), it's a heading
            if common_count < word_count / 2:
                return True
    
    # Pattern 4: Title Case (2-12 words, 50%+ capitalized, no sentence punctuation)
    if 2 <= word_count <= 12:
        capitalized = sum(1 for w in words if w and w[0].isupper())
        cap_ratio = capitalized / word_count
        
        if cap_ratio >= 0.5 and not line.endswith(('.', ',', ';', '!', '?')):
            # Exclude if contains too many common sentence words
            lower_words = [w.lower() for w in words]
            sentence_words = {'the', 'a', 'an', 'is', 'are', 'was', 'were', 'to', 'of', 'in', 'on', 'for', 'with', 'at', 'by'}
            common_count = len(sentence_words & set(lower_words))
            
            # Allow if few sentence words
            if common_count <= 2:
                return True
    
    return False


def normalize_section_name(name: str) -> str:
    """
    Convert section name to snake_case for consistent matching.
    Example: "Executive Summary" -> "executive_summary"
    """
    # Remove numbering (e.g., "1.2 Methodology" -> "Methodology")
    name = re.sub(r'^\d+(\.\d+)*\.?\s*', '', name)
    
    # Convert to lowercase and snake_case
    name = name.lower()
    name = re.sub(r'[^\w\s-]', '', name)
    name = re.sub(r'[-\s]+', '_', name)
    
    return name.strip('_')


def denormalize_section_name(key: str) -> str:
    """
    Convert snake_case back to Title Case for display.
    Example: "executive_summary" -> "Executive Summary"
    """
    return key.replace('_', ' ').title()




