import logging
import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from io import BytesIO
from fpdf import FPDF


logger = logging.getLogger(__name__)


def generate_json_filename(base="extracted_report", extension="json"):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{base}_{timestamp}.{extension}"

def save_report(data, output_dir="outputs"):
    os.makedirs(output_dir, exist_ok=True)
    filename = generate_json_filename()
    path = os.path.join(output_dir, filename)
    with open(path, "w", encoding="utf-8") as f:
        import json
        json.dump(data, f, indent=2, ensure_ascii=False)
    return path

def add_bold_runs(paragraph, text):
        """
        Add runs to a docx paragraph, making text between ** ** bold.
        """
        parts = re.split(r'(\*\*.+?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)

def parse_list_lines(text):
    lines = text.strip().split("\n")
    parsed = []
    for line in lines:
        original = line
        line = line.rstrip()
        leading_spaces = len(line) - len(line.lstrip())
        stripped = line.lstrip()

        if re.match(r"^\d+\.", stripped) and leading_spaces == 0:
            parsed.append(("level1", stripped))
        elif stripped.startswith("-") and 1 <= leading_spaces <= 4:
            parsed.append(("level2", stripped))
        else:
            parsed.append(("text", stripped))
    return parsed

def is_new_section(line: str) -> bool:
    """
    Define what constitutes a new section (e.g., headings like '1.1 Scope of Work').
    """
    return bool(re.match(r'^\d+\.\d+\s+', line.strip()))  # Matches "1.1", "2.3", etc.

def add_nested_list_paragraphs(doc, text):
    parsed_lines = parse_list_lines(text)

    for i, (level, line) in enumerate(parsed_lines):
        if level == "level1":
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.25)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            add_bold_runs(para, line)
        elif level == "level2":
            line_text = re.sub(r'^-\s*', '', line)
            para = doc.add_paragraph(style="List Bullet 3")
            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            add_bold_runs(para, line_text)
        else:
            para = doc.add_paragraph()
            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            add_bold_runs(para, line)

def remove_leading_list_marker(text):
    """Removes leading list symbols like '1.', '-', or 'a.'."""
    return re.sub(r'^(\d+\.|[a-zA-Z]\.|-)\s*', '', text).strip()

def is_list_item(text):
    """Detects if a line starts with a number, letter, or dash followed by a space."""
    return re.match(r'^(\d+\.|[a-zA-Z]\.|-)\s+', text.strip()) is not None

def set_document_margins(doc, top=1, bottom=1, left=1, right=1):
    """
    Sets the margins of the Word document.

    Args:
        doc (Document): A python-docx Document object.
        top (float): Top margin in inches.
        bottom (float): Bottom margin in inches.
        left (float): Left margin in inches.
        right (float): Right margin in inches.
    """
    section = doc.sections[0]
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)

def set_document_font(doc, font_name="Arial"):
    """
    Sets the default font for the Word document, including East Asian font settings.

    Args:
        doc (Document): A python-docx Document object.
        font_name (str): The name of the font to apply.
    """
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    rFonts = font.element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), font_name)

def add_styled_heading(doc, text, level=1, font_size=None, space_before=20, space_after=10):
    """
    Adds a styled heading with level-based default font size, unless manually overridden.

    Args:
        doc (Document): The Word document object.
        text (str): The heading text.
        level (int): The heading level (1 = Heading 1, 2 = Heading 2, etc.).
        font_size (int or float, optional): Override font size in points. Defaults to 24 if level 1, 16 otherwise.
        space_before (int or float): Space before the heading in points.
        space_after (int or float): Space after the heading in points.
    """
    heading_para = doc.add_heading(text, level=level)
    run = heading_para.runs[0]

    # Apply dynamic font size based on heading level unless overridden
    if font_size is None:
        font_size = 24 if level == 1 else 16
    run.font.size = Pt(font_size)

    # Set spacing before and after
    heading_para.paragraph_format.space_before = Pt(space_before)
    heading_para.paragraph_format.space_after = Pt(space_after)

def write_section(doc, title, content, mode, level=1):
    if title.strip():
        add_styled_heading(doc, title.strip(), level=level)

    content = content.replace("TERMINATE", "").strip()

    if mode == "extracted" and not content:
        doc.add_paragraph("No relevant information found.", style="Intense Quote")
        return

    for paragraph in re.split(r'\n\s*\n', content):
        stripped_paragraph = paragraph.strip()
        if not stripped_paragraph:
            continue

        # Handle inline markdown-style headers
        if mode == "final":
            if stripped_paragraph.startswith("# "):
                add_styled_heading(doc, stripped_paragraph[2:].strip(), level=1)
                continue
            elif stripped_paragraph.startswith("## "):
                add_styled_heading(doc, stripped_paragraph[3:].strip(), level=2)
                continue

        if any(is_list_item(line) for line in stripped_paragraph.splitlines()):
            add_nested_list_paragraphs(doc, stripped_paragraph)
        else:
            para = doc.add_paragraph()
            add_bold_runs(para, stripped_paragraph)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

def write_to_docx(section_outputs, filename="output.docx", mode="draft"):
    doc = Document()
    set_document_margins(doc) # Set up margins
    set_document_font(doc) # Set up font

    for section_key, section_data in section_outputs.items():
        title = section_data.get("title", section_key)
        content = section_data.get("content") or section_data.get("text", "")

        if "subsections" in section_data:
            write_section(doc, title, "", mode, level=1)
            for sub_key, sub_data in section_data["subsections"].items():
                sub_title = sub_data.get("title", sub_key)
                sub_content = sub_data.get("content", "")
                logger.debug("Writing subsection content: %s\n%s", sub_title, sub_content)
                write_section(doc, sub_title, sub_content, mode, level=2)
        else:
            logger.debug("Writing section content: %s\n%s", title, content);
            write_section(doc, title, content, mode, level=1)

    doc.save(filename)
